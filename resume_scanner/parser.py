"""
Resume Parser Module - Production-Grade Edition
Handles robust multi-format parsing (PDF, DOCX, TXT), structural decomposition,
scanned PDF detection, section identification, and contact information extraction.
"""

import re
import io
from pathlib import Path
from typing import Optional, Dict, List, Tuple, Any
from dataclasses import dataclass, field

from .config import ParserConfig, DEFAULT_PARSER_CONFIG


@dataclass
class ResumeDocument:
    """
    Structured internal representation of a parsed resume document.
    Additive data model preserving pages, paragraphs, sections, and bullets.
    """

    raw_text: str = ""
    clean_text: str = ""
    pages: List[str] = field(default_factory=list)
    sections: Dict[str, str] = field(default_factory=dict)
    section_presence: Dict[str, bool] = field(default_factory=dict)
    paragraphs: List[str] = field(default_factory=list)
    bullet_points: List[str] = field(default_factory=list)
    contact_info: Dict[str, Optional[str]] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)

    @property
    def page_count(self) -> int:
        return len(self.pages) if self.pages else 1

    @property
    def is_scanned(self) -> bool:
        return self.metadata.get("is_scanned", False)


class ResumeParser:
    """
    Production-grade resume parser supporting PDF, DOCX, and TXT formats.
    Provides backward-compatible text output and structured ResumeDocument extraction.
    """

    def __init__(self, config: Optional[ParserConfig] = None):
        self.config = config or DEFAULT_PARSER_CONFIG
        self.text: str = ""
        self.metadata: Dict[str, Any] = {}
        self.document: Optional[ResumeDocument] = None

    # ─────────────────────────────────────────────────────────────────────────
    # Public API: string parse (backward-compatible)
    # ─────────────────────────────────────────────────────────────────────────
    def parse(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        file_type: Optional[str] = None,
    ) -> str:
        """
        Parse a resume file and extract cleaned text.
        Maintains full backward compatibility with the existing public interface.
        """
        doc = self.parse_document(
            file_path=file_path, file_content=file_content, file_type=file_type
        )
        self.text = doc.clean_text
        self.metadata = doc.metadata
        return self.text

    # ─────────────────────────────────────────────────────────────────────────
    # Public API: structured ResumeDocument
    # ─────────────────────────────────────────────────────────────────────────
    def parse_document(
        self,
        file_path: Optional[str] = None,
        file_content: Optional[bytes] = None,
        file_type: Optional[str] = None,
    ) -> ResumeDocument:
        """
        Parse resume into a rich, structured ResumeDocument.
        """
        if file_content is not None and file_type:
            raw_text, pages, meta = self._parse_from_bytes(file_content, file_type)
        elif file_path:
            raw_text, pages, meta = self._parse_from_path(file_path)
        else:
            raise ValueError(
                "Either file_path or (file_content and file_type) must be provided"
            )

        clean_text = self._clean_text(raw_text)
        self.text = clean_text
        self.metadata = meta

        # Minimum text length validation
        if len(clean_text.strip()) < self.config.min_extracted_text_length:
            if meta.get("is_scanned", False):
                meta["warning"] = (
                    "Scanned or image-based document detected with minimal selectable text. "
                    "For accurate ATS analysis, please upload a text-based PDF, DOCX, or TXT file."
                )
            else:
                meta["warning"] = (
                    "Document contains very little or no extracted text. "
                    "Please verify the file contains selectable text and is not empty."
                )

        # Structural decomposition
        paragraphs = self.extract_paragraphs(clean_text)
        bullet_points = self.extract_bullet_points(clean_text)
        section_texts = self.get_section_content(clean_text)
        section_flags = self.get_sections(clean_text)
        contact_info = self.extract_contact_info(clean_text)

        self.document = ResumeDocument(
            raw_text=raw_text,
            clean_text=clean_text,
            pages=pages,
            sections=section_texts,
            section_presence=section_flags,
            paragraphs=paragraphs,
            bullet_points=bullet_points,
            contact_info=contact_info,
            metadata=meta,
        )

        return self.document

    # ─────────────────────────────────────────────────────────────────────────
    # File handling and format dispatch
    # ─────────────────────────────────────────────────────────────────────────
    def _parse_from_path(self, file_path: str) -> Tuple[str, List[str], Dict[str, Any]]:
        path = Path(file_path)
        if not path.is_file():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()
        file_size = path.stat().st_size

        if file_size > self.config.max_file_size_bytes:
            raise ValueError(
                f"File size ({file_size / (1024*1024):.1f} MB) exceeds maximum allowed "
                f"size of {self.config.max_file_size_mb} MB"
            )

        if file_size == 0:
            raise ValueError("The provided file is empty (0 bytes).")

        with open(path, "rb") as f:
            content = f.read()

        return self._parse_from_bytes(content, suffix, filename=path.name)

    def _parse_from_bytes(
        self, content: bytes, file_type: str, filename: str = "upload"
    ) -> Tuple[str, List[str], Dict[str, Any]]:
        file_type = file_type.lower()
        if not file_type.startswith("."):
            file_type = "." + file_type

        # Reject empty content
        if not content or len(content) == 0:
            raise ValueError("The uploaded file contains no data (0 bytes).")

        # Check maximum file size
        if len(content) > self.config.max_file_size_bytes:
            raise ValueError(
                f"Uploaded file exceeds the maximum limit of {self.config.max_file_size_mb} MB."
            )

        # Explicitly handle legacy binary .doc
        if file_type in self.config.unsupported_legacy_extensions:
            raise ValueError(
                "Legacy binary .doc files are not supported. "
                "Please save your file as modern .docx or .pdf in Word/Google Docs and re-upload."
            )

        if file_type not in self.config.supported_extensions:
            supported = ", ".join(self.config.supported_extensions)
            raise ValueError(f"Unsupported file format '{file_type}'. Supported formats: {supported}")

        meta: Dict[str, Any] = {
            "filename": filename,
            "file_type": file_type,
            "file_size_bytes": len(content),
            "is_scanned": False,
        }

        pages: List[str] = []
        raw_text: str = ""

        if file_type == ".pdf":
            raw_text, pages, is_scanned = self._parse_pdf(content)
            meta["is_scanned"] = is_scanned
            meta["page_count"] = len(pages)
        elif file_type == ".docx":
            raw_text, paragraphs = self._parse_docx(content)
            pages = [raw_text]
            meta["paragraph_count"] = len(paragraphs)
        elif file_type == ".txt":
            raw_text = self._parse_txt(content)
            pages = [raw_text]

        return raw_text, pages, meta

    # ─────────────────────────────────────────────────────────────────────────
    # Format-specific extractors
    # ─────────────────────────────────────────────────────────────────────────
    def _parse_pdf(self, content: bytes) -> Tuple[str, List[str], bool]:
        """Extract text page-by-page using PyMuPDF (fitz) with scanned document detection."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF (fitz) is required for PDF parsing. Install with: pip install PyMuPDF"
            )

        pages: List[str] = []
        is_scanned = False
        has_images = False

        try:
            pdf_document = fitz.open(stream=content, filetype="pdf")
        except Exception as e:
            raise ValueError(f"Failed to open PDF document: {e}. The file may be corrupt or encrypted.")

        try:
            total_pages = len(pdf_document)
            if total_pages == 0:
                pdf_document.close()
                return "", [], False

            total_chars = 0
            for page_num in range(total_pages):
                page = pdf_document[page_num]
                page_text = page.get_text() or ""
                pages.append(page_text)
                total_chars += len(page_text.strip())

                # Check if page has embedded images
                if not has_images:
                    image_list = page.get_images()
                    if image_list:
                        has_images = True

            pdf_document.close()

            # Scanned detection: images present but negligible selectable text
            avg_chars_per_page = total_chars / max(total_pages, 1)
            if has_images and avg_chars_per_page < self.config.scanned_pdf_char_threshold_per_page:
                is_scanned = True

            raw_text = "\n\n".join(pages)
            return raw_text, pages, is_scanned

        except Exception as e:
            try:
                pdf_document.close()
            except Exception:
                pass
            raise ValueError(f"Error extracting content from PDF: {str(e)}")

    def _parse_docx(self, content: bytes) -> Tuple[str, List[str]]:
        """Extract text from DOCX preserving paragraph order and tables."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install with: pip install python-docx"
            )

        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {e}. The file may be corrupt or an unsupported format.")

        text_parts: List[str] = []

        # Extract regular paragraphs
        for paragraph in doc.paragraphs:
            p_text = paragraph.text.strip()
            if p_text:
                text_parts.append(p_text)

        # Extract tables in order
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    # Deduplicate adjacent identical cells from merged tables
                    unique_cells = []
                    for c in row_cells:
                        if not unique_cells or c != unique_cells[-1]:
                            unique_cells.append(c)
                    text_parts.append(" | ".join(unique_cells))

        raw_text = "\n".join(text_parts)
        return raw_text, text_parts

    def _parse_txt(self, content: bytes) -> str:
        """Decode TXT bytes with multi-encoding fallback."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
        for enc in encodings:
            try:
                return content.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return content.decode("utf-8", errors="replace")

    # ─────────────────────────────────────────────────────────────────────────
    # Text normalization: preserves line breaks & structural boundaries
    # ─────────────────────────────────────────────────────────────────────────
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text without destroying paragraph/line breaks.
        """
        if not text:
            return ""

        # Normalize carriage returns
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Replace non-breaking spaces and zero-width characters with regular spaces
        text = re.sub(r"[\u00a0\u1680\u2000-\u200b\u202f\u205f\u3000\ufeff]", " ", text)

        # Clean horizontal whitespace per line (tabs, multiple spaces)
        lines = []
        for line in text.split("\n"):
            # Collapse multiple spaces on a single line
            clean_line = re.sub(r"[ \t]+", " ", line).strip()
            lines.append(clean_line)

        # Join back preserving single line breaks
        text = "\n".join(lines)

        # Collapse more than 2 consecutive newlines into 2
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    # ─────────────────────────────────────────────────────────────────────────
    # Structural decomposition
    # ─────────────────────────────────────────────────────────────────────────
    def extract_paragraphs(self, text: Optional[str] = None) -> List[str]:
        """Extract cohesive paragraphs separated by double newlines."""
        t = text or self.text
        if not t:
            return []
        raw_paras = re.split(r"\n\s*\n+", t)
        return [p.strip() for p in raw_paras if len(p.strip()) > 15]

    def extract_bullet_points(self, text: Optional[str] = None) -> List[str]:
        """Extract individual bullet points from resume text."""
        t = text or self.text
        if not t:
            return []

        bullets = []
        bullet_prefix_regex = r"^\s*([•○■►\*\-\–\—]|\d+\.|\([a-z0-9]+\))\s+(.*)$"

        for line in t.split("\n"):
            line_str = line.strip()
            if not line_str:
                continue
            match = re.match(bullet_prefix_regex, line_str)
            if match:
                content = match.group(2).strip()
                if len(content) > 5:
                    bullets.append(content)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                bullets.append(line_str[2:].strip())

        return bullets

    # ─────────────────────────────────────────────────────────────────────────
    # Context-aware section extraction
    # ─────────────────────────────────────────────────────────────────────────
    SECTION_HEADER_PATTERNS = {
        "contact": r"^\s*(?:#+\s*)?(?:contact(?:\s+information)?|personal\s+details)\s*:?\s*$",
        "summary": r"^\s*(?:#+\s*)?(?:professional\s+summary|summary|executive\s+summary|profile|about\s+me|career\s+objective)\s*:?\s*$",
        "experience": r"^\s*(?:#+\s*)?(?:experience|work\s+experience|professional\s+experience|employment\s+history|career\s+history|work\s+history)\s*:?\s*$",
        "education": r"^\s*(?:#+\s*)?(?:education|academic\s+background|qualifications|academic\s+qualifications|degrees?)\s*:?\s*$",
        "skills": r"^\s*(?:#+\s*)?(?:skills|technical\s+skills|core\s+competencies|technical\s+expertise|skills\s+(?:&|and)\s+technologies|technologies)\s*:?\s*$",
        "projects": r"^\s*(?:#+\s*)?(?:projects|key\s+projects|selected\s+projects|personal\s+projects|portfolio)\s*:?\s*$",
        "certifications": r"^\s*(?:#+\s*)?(?:certifications?|certificates?|licenses?(?:\s+&\s+certifications?)?|credentials)\s*:?\s*$",
    }

    def get_sections(self, text: Optional[str] = None) -> Dict[str, bool]:
        """
        Identify section presence.
        Returns a dictionary mapping section keys to a boolean (backward compatible).
        """
        content_dict = self.get_section_content(text)
        return {key: bool(content.strip()) for key, content in content_dict.items()}

    def get_section_content(
        self,
        section_name: Optional[str] = None,
        text: Optional[str] = None,
    ) -> Any:
        """
        Extract the text content associated with resume sections.
        If section_name is provided (e.g. 'experience'), returns that section's text.
        If omitted or if full text is passed, returns Dict[str, str] of all sections.
        """
        if section_name and section_name in self.SECTION_HEADER_PATTERNS:
            t = text or self.text
            all_sec = self._parse_all_sections(t)
            return all_sec.get(section_name, "")

        t = text or section_name or self.text
        return self._parse_all_sections(t)

    def _parse_all_sections(self, t: Optional[str]) -> Dict[str, str]:
        if not t:
            return {k: "" for k in self.SECTION_HEADER_PATTERNS}

        lines = t.split("\n")
        section_starts: List[Tuple[int, str]] = []

        for idx, line in enumerate(lines):
            clean_line = line.strip()
            if not clean_line or len(clean_line) > 55:
                continue

            for s_name, pattern in self.SECTION_HEADER_PATTERNS.items():
                if re.match(pattern, clean_line, re.IGNORECASE):
                    section_starts.append((idx, s_name))
                    break

        sections: Dict[str, str] = {k: "" for k in self.SECTION_HEADER_PATTERNS}

        if not section_starts:
            lower_t = t.lower()
            for s_name, s_pat in self.SECTION_HEADER_PATTERNS.items():
                if re.search(s_pat, lower_t, re.MULTILINE | re.IGNORECASE):
                    sections[s_name] = "Found via pattern"
            return sections

        for i, (line_idx, sec_name) in enumerate(section_starts):
            start = line_idx + 1
            end = section_starts[i + 1][0] if i + 1 < len(section_starts) else len(lines)
            sec_text = "\n".join(lines[start:end]).strip()
            sections[sec_name] = sec_text

        return sections

    # ─────────────────────────────────────────────────────────────────────────
    # Contact information extraction
    # ─────────────────────────────────────────────────────────────────────────
    def extract_contact_info(self, text: Optional[str] = None) -> Dict[str, Optional[str]]:
        """
        Extract normalized email, phone, LinkedIn, GitHub, and portfolio links.
        Avoids false positives on dates, zip codes, and numerical ranges.
        """
        t = text or self.text
        contact: Dict[str, Optional[str]] = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "website": None,
        }

        if not t:
            return contact

        # 1. Email (RFC-compliant standard pattern)
        email_match = re.search(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b", t
        )
        if email_match:
            contact["email"] = email_match.group(0).strip()

        # 2. Phone (International and local formats with strict digit count: 7-15 digits)
        # Excludes raw year ranges like 2020-2023 or 2021 - 2024
        phone_patterns = [
            r"(?:\+?\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b",
            r"\+\d{1,3}[\s.-]?\d{2,4}[\s.-]?\d{3,4}[\s.-]?\d{3,5}\b",
        ]
        for pat in phone_patterns:
            phone_matches = re.finditer(pat, t)
            for pm in phone_matches:
                candidate = pm.group(0).strip()
                digits_only = re.sub(r"\D", "", candidate)
                # Avoid dates: 1990-2026
                if len(digits_only) >= 7 and len(digits_only) <= 15:
                    if not (len(digits_only) == 8 and digits_only.startswith("20")):
                        contact["phone"] = candidate
                        break
            if contact["phone"]:
                break

        # 3. LinkedIn
        li_match = re.search(
            r"(?:https?://)?(?:www\.)?linkedin\.com/in/([a-zA-Z0-9_-]+)/?",
            t,
            re.IGNORECASE,
        )
        if li_match:
            contact["linkedin"] = f"https://linkedin.com/in/{li_match.group(1)}"

        # 4. GitHub
        gh_match = re.search(
            r"(?:https?://)?(?:www\.)?github\.com/([a-zA-Z0-9_-]+)/?",
            t,
            re.IGNORECASE,
        )
        if gh_match:
            username = gh_match.group(1)
            # Exclude generic github paths
            if username.lower() not in ["features", "explore", "topics", "collections"]:
                contact["github"] = f"https://github.com/{username}"

        # 5. Website / Portfolio
        web_match = re.search(
            r"(?:portfolio|website|web)\s*:\s*(https?://[^\s]+|[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?)",
            t,
            re.IGNORECASE,
        )
        if web_match:
            raw_url = web_match.group(1).strip()
            if not raw_url.startswith("http"):
                raw_url = f"https://{raw_url}"
            contact["website"] = raw_url
        else:
            # Look for dev/io/me domains
            custom_domain = re.search(
                r"\b(?:https?://)?([a-zA-Z0-9-]+\.(?:dev|io|me|tech|site|online))\b",
                t,
                re.IGNORECASE,
            )
            if custom_domain:
                contact["website"] = f"https://{custom_domain.group(1)}"

        return contact
